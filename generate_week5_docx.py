"""
==============================================
WEEK 5 DOCUMENTATION GENERATOR - DOCX FORMAT
==============================================
Generates a professionally formatted Word document
Double-spaced, Times New Roman font
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path


def set_font(run, size=12, bold=False, italic=False, color=None):
    """Set font properties for a run"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(12)
    heading.paragraph_format.line_spacing = 2.0
    
    run = heading.add_run(text)
    size = 18 if level == 1 else 14 if level == 2 else 12
    bold = level <= 2
    set_font(run, size=size, bold=bold)
    
    if level == 1:
        heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with double spacing"""
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(6)
    
    for run in p.runs:
        set_font(run, size=12, bold=bold, italic=italic)
    
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Inches(0.5 + (level * 0.25))
    
    for run in p.runs:
        set_font(run, size=12)


def create_week5_document():
    """Create Week 5 completion document"""
    doc = Document()
    
    # Set default font for document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title Page
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(72)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.line_spacing = 2.0
    
    run = title.add_run('WEEK 5 COMPLETION REPORT')
    set_font(run, size=18, bold=True)
    
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.line_spacing = 2.0
    
    run = subtitle.add_run('Data and Environment Hardening')
    set_font(run, size=14, bold=True)
    
    # Metadata
    metadata_items = [
        'Status: ✅ COMPLETE',
        'Date Completed: July 1, 2026',
        'Focus Area: Data and Environment Hardening',
        'System: Secure File Transfer System'
    ]
    
    for item in metadata_items:
        add_paragraph(doc, item)
    
    doc.add_paragraph()  # Spacing
    
    # Executive Summary
    add_heading(doc, 'Executive Summary', level=1)
    add_paragraph(doc, 
        'Week 5 successfully implemented enterprise-grade environment configuration and secrets '
        'management for the Secure File Transfer System. All hardcoded configuration has been '
        'moved to environment variables, comprehensive validation was added, and the system is '
        'now deployment-ready with clear separation between code and configuration.')
    
    # Deliverables
    add_heading(doc, 'Deliverables Completed', level=1)
    
    add_heading(doc, '1. Configuration Management System', level=2)
    add_bullet(doc, 'File: config.py')
    add_bullet(doc, 'Loads configuration from .env file using python-dotenv', level=1)
    add_bullet(doc, 'Centralized access to all app settings', level=1)
    add_bullet(doc, 'Type-safe configuration (int, bool, Path, list)', level=1)
    add_bullet(doc, 'Comprehensive validation with detailed error messages', level=1)
    add_bullet(doc, 'Graceful fallbacks for missing values', level=1)
    add_bullet(doc, 'Auto-creates required directories', level=1)
    
    add_heading(doc, '2. Environment Configuration Files', level=2)
    add_bullet(doc, 'File: .env.example – Template with all configuration options')
    add_bullet(doc, 'File: .env – Pre-configured local development settings')
    
    add_heading(doc, '3. Secrets Removed from Code', level=2)
    add_paragraph(doc, 'Before Week 5:')
    add_bullet(doc, 'Hardcoded SECRET_KEY in app/__init__.py')
    add_bullet(doc, 'Hardcoded demo credentials in auth.py')
    add_bullet(doc, 'Hardcoded paths and ports')
    
    add_paragraph(doc, 'After Week 5:')
    add_bullet(doc, 'SECRET_KEY from environment variable')
    add_bullet(doc, 'Demo users configurable via DEMO_USERS')
    add_bullet(doc, 'All paths and ports configurable')
    
    add_heading(doc, '4. Startup Validation System', level=2)
    add_bullet(doc, 'File: validate_startup.py')
    add_bullet(doc, 'Configuration file existence check', level=1)
    add_bullet(doc, 'Configuration loading verification', level=1)
    add_bullet(doc, 'Required directories validation', level=1)
    add_bullet(doc, 'Python dependencies verification', level=1)
    add_bullet(doc, 'Demo users configuration check', level=1)
    add_bullet(doc, 'Security settings audit', level=1)
    
    add_heading(doc, '5. Application Integration', level=2)
    add_bullet(doc, 'app/__init__.py – Updated to use config module')
    add_bullet(doc, 'run.py – Displays configuration on startup')
    add_bullet(doc, 'app/modules/auth.py – Uses configured demo users')
    add_bullet(doc, 'app/modules/logger.py – Uses configured log folder')
    
    add_heading(doc, '6. Storage Mode Decision', level=2)
    add_paragraph(doc, 'Decision: In-Memory Storage (Academic Prototype)')
    add_paragraph(doc, 'Rationale: Simplicity for academic context, focus on encryption and metrics')
    add_paragraph(doc, 'Configuration: STORAGE_MODE=in-memory')
    add_paragraph(doc, 'Future: Path documented for SQLite migration')
    
    add_heading(doc, '7. Documentation', level=2)
    add_bullet(doc, 'WEEK5_ENVIRONMENT_HARDENING.md – Technical reference (400+ lines)')
    add_bullet(doc, 'WEEK5_COMPLETION_REPORT.md – Executive summary')
    add_bullet(doc, 'ENVIRONMENT_SETUP_GUIDE.md – Step-by-step setup instructions')
    add_bullet(doc, 'WEEK5_STATUS_UPDATE.md – Project tracking')
    
    doc.add_page_break()
    
    # Configuration Coverage
    add_heading(doc, 'Configuration Coverage', level=1)
    
    coverage_items = [
        ('Flask', 'SECRET_KEY, ENV, DEBUG, HOST, PORT'),
        ('File Upload', 'MAX_FILE_SIZE_MB, UPLOAD_FOLDER, subfolder paths'),
        ('Logging', 'LOG_FOLDER, LOG_LEVEL'),
        ('Security', 'SESSION_TIMEOUT_MINUTES, MAX_LOGIN_ATTEMPTS, LOCKOUT_DURATION_MINUTES'),
        ('Encryption', 'ENCRYPTION_ALGORITHM, KEY_ROTATION_POLICY'),
        ('Demo Users', 'Configurable list (format: username:password:email)'),
        ('Storage', 'STORAGE_MODE (in-memory or database)'),
        ('Research', 'METRICS_ENABLED, METRICS_OUTPUT_FILE')
    ]
    
    add_paragraph(doc, 'All Configurable Settings:', bold=True)
    for category, settings in coverage_items:
        add_bullet(doc, f'{category}: {settings}')
    
    # Security Improvements
    add_heading(doc, 'Security Improvements', level=1)
    
    add_paragraph(doc, 'Before Week 5:', bold=True)
    add_bullet(doc, 'Hardcoded SECRET_KEY in Python')
    add_bullet(doc, 'Hardcoded demo credentials')
    add_bullet(doc, 'Hardcoded paths and configuration')
    add_bullet(doc, 'No startup validation')
    add_bullet(doc, 'No environment-specific settings')
    
    add_paragraph(doc, 'After Week 5:', bold=True)
    add_bullet(doc, 'SECRET_KEY from environment')
    add_bullet(doc, 'All credentials in .env')
    add_bullet(doc, 'All paths configurable')
    add_bullet(doc, '6-check startup validation')
    add_bullet(doc, 'Dev/Test/Prod ready')
    add_bullet(doc, 'Security warnings for misconfigurations')
    
    # Files Created and Modified
    add_heading(doc, 'Files Created and Modified', level=1)
    
    add_paragraph(doc, 'Files Created:', bold=True)
    files_created = [
        ('config.py', '~280 lines'),
        ('.env', '~45 lines'),
        ('.env.example', '~75 lines'),
        ('validate_startup.py', '~200 lines'),
        ('WEEK5_ENVIRONMENT_HARDENING.md', '~400 lines')
    ]
    
    for filename, size in files_created:
        add_bullet(doc, f'{filename} ({size})')
    
    add_paragraph(doc, 'Files Modified:', bold=True)
    files_modified = [
        'app/__init__.py',
        'run.py',
        'app/modules/auth.py',
        'app/modules/logger.py'
    ]
    
    for filename in files_modified:
        add_bullet(doc, filename)
    
    doc.add_page_break()
    
    # How to Use
    add_heading(doc, 'How to Use', level=1)
    
    add_heading(doc, 'Quick Start', level=2)
    add_bullet(doc, 'Run: python validate_startup.py')
    add_bullet(doc, 'Run: python run.py')
    add_bullet(doc, 'Access application at http://127.0.0.1:5000')
    
    add_heading(doc, 'Change Configuration', level=2)
    add_bullet(doc, 'Edit .env file')
    add_bullet(doc, 'Change desired settings')
    add_bullet(doc, 'Restart application')
    
    add_heading(doc, 'Deployment Setup', level=2)
    add_bullet(doc, 'Copy .env.example to .env')
    add_bullet(doc, 'Set FLASK_SECRET_KEY to secure random value')
    add_bullet(doc, 'Set FLASK_ENV=production')
    add_bullet(doc, 'Set DEBUG=False')
    add_bullet(doc, 'Run validation and start')
    
    # Configuration Reference
    add_heading(doc, 'Configuration Reference', level=1)
    
    add_heading(doc, 'Core Settings', level=2)
    config_ref = [
        'FLASK_SECRET_KEY – Session encryption key (change for production)',
        'FLASK_ENV – Environment: development, testing, production',
        'DEBUG – Enable Flask debug mode (False for production)',
        'FLASK_HOST – Server host (127.0.0.1 for local, 0.0.0.0 for network)',
        'FLASK_PORT – Server port (default 5000)'
    ]
    
    for item in config_ref:
        add_bullet(doc, item)
    
    add_heading(doc, 'File Management', level=2)
    file_ref = [
        'MAX_FILE_SIZE_MB – Maximum file upload size (default 50 MB)',
        'UPLOAD_FOLDER – Base upload directory',
        'PLAIN_UPLOAD_FOLDER – Unencrypted files directory',
        'ENCRYPTED_UPLOAD_FOLDER – Encrypted files directory'
    ]
    
    for item in file_ref:
        add_bullet(doc, item)
    
    add_heading(doc, 'Logging', level=2)
    log_ref = [
        'LOG_FOLDER – Directory for logs',
        'LOG_LEVEL – Logging level (DEBUG, INFO, WARNING, ERROR, CRITICAL)'
    ]
    
    for item in log_ref:
        add_bullet(doc, item)
    
    add_heading(doc, 'Security', level=2)
    sec_ref = [
        'SESSION_TIMEOUT_MINUTES – Session duration (default 30)',
        'MAX_LOGIN_ATTEMPTS – Failed attempts before lockout (default 5)',
        'LOCKOUT_DURATION_MINUTES – Lockout window duration (default 15)'
    ]
    
    for item in sec_ref:
        add_bullet(doc, item)
    
    doc.add_page_break()
    
    # Testing Completed
    add_heading(doc, 'Testing Completed', level=1)
    
    tests = [
        'Config loads without errors',
        'Configuration displays on startup',
        'All paths resolve correctly',
        'Demo users load from config',
        'Environment variables override defaults',
        'Directory creation automatic',
        'Validation script works correctly',
        'Security checks function properly',
        'Error messages are clear and helpful'
    ]
    
    for test in tests:
        add_bullet(doc, test)
    
    # Success Metrics
    add_heading(doc, 'Success Metrics', level=1)
    
    metrics = [
        ('Configuration modules created', '1', '✅'),
        ('Hardcoded values removed', '100%', '✅'),
        ('Configuration options', '20+', '22 ✅'),
        ('Validation checks', '5+', '6 ✅'),
        ('Documentation lines', '300+', '400+ ✅'),
        ('Files modified', '4', '4 ✅'),
        ('Deployment readiness', 'Yes', 'Yes ✅')
    ]
    
    for metric, target, achieved in metrics:
        add_paragraph(doc, f'{metric}: Target {target} – Achieved {achieved}')
    
    # Week 6 Ready
    add_heading(doc, 'Ready for Week 6', level=1)
    add_paragraph(doc, 
        'The environment configuration is now complete and ready to support Week 6 work on '
        'Authentication Completion. All configuration settings for session management, login '
        'attempts, and account lockout are prepared and ready for implementation.')
    
    add_paragraph(doc, 'Configuration prepared for Week 6:', bold=True)
    add_bullet(doc, 'SESSION_TIMEOUT_MINUTES=30')
    add_bullet(doc, 'MAX_LOGIN_ATTEMPTS=5')
    add_bullet(doc, 'LOCKOUT_DURATION_MINUTES=15')
    
    doc.add_page_break()
    
    # Summary
    add_heading(doc, 'Summary', level=1)
    add_paragraph(doc, 
        'Week 5 successfully implemented enterprise-grade environment management. The system '
        'now features centralized configuration, no secrets in code, comprehensive validation, '
        'and is deployment-ready. All objectives have been achieved and the system is prepared '
        'for continued development in Week 6.')
    
    add_heading(doc, 'Key Achievements', level=2)
    achievements = [
        'Centralized configuration system implemented',
        'All hardcoded values moved to environment',
        'Comprehensive startup validation system',
        'Deployment-ready configuration',
        'Research reproducibility ensured',
        'Security hardened and documented',
        'Clear path for future database migration'
    ]
    
    for achievement in achievements:
        add_bullet(doc, achievement)
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Week 5 Complete: 100% of objectives achieved', bold=True)
    
    return doc


def main():
    """Generate and save the document"""
    print("Generating Week 5 DOCX document...")
    
    doc = create_week5_document()
    
    output_path = Path(__file__).parent / 'WEEK5_DOCUMENTATION.docx'
    doc.save(str(output_path))
    
    print(f"✓ Document saved: {output_path}")
    print(f"✓ Format: DOCX (Microsoft Word)")
    print(f"✓ Spacing: Double-spaced")
    print(f"✓ Font: Times New Roman, 12pt")
    print(f"✓ Total pages: Approximately 10-12")


if __name__ == '__main__':
    main()
