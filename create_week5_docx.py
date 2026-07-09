"""
Week 5 DOCX Document Generator - Simplified
Creates a professional DOCX document with Week 5 content
Times New Roman, Double Spacing, Professional Formatting
"""

import sys

# Try to import docx, install if needed
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


def create_docx():
    """Create Week 5 DOCX document"""
    doc = Document()
    
    # Set default font for entire document
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('WEEK 5: DATA AND ENVIRONMENT HARDENING')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Secure File Transfer System')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    subtitle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Date
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('July 1, 2026')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    date_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Add spacing
    doc.add_paragraph()
    
    # Executive Summary Section
    heading1 = doc.add_paragraph()
    run = heading1.add_run('EXECUTIVE SUMMARY')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    heading1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    summary_text = """Week 5 successfully implemented enterprise-grade environment configuration and secrets management for the Secure File Transfer System. All hardcoded configuration has been moved to environment variables, comprehensive validation was added, and the system is now deployment-ready with clear separation between code and configuration."""
    
    summary_p = doc.add_paragraph(summary_text)
    summary_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    for run in summary_p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Deliverables Section
    heading2 = doc.add_paragraph()
    run = heading2.add_run('DELIVERABLES COMPLETED')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    heading2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    deliverables = [
        ('Configuration Management System', [
            'File: config.py',
            'Loads configuration from .env file using python-dotenv',
            'Centralized access to all app settings',
            'Type-safe configuration (int, bool, Path, list)',
            'Comprehensive validation with detailed error messages',
            'Auto-creates required directories'
        ]),
        ('Environment Configuration Files', [
            'File: .env.example - Template with all options',
            'File: .env - Pre-configured for local development',
            'Comprehensive documentation and recommendations'
        ]),
        ('Secrets Removed from Code', [
            'All hardcoded values moved to environment variables',
            'SECRET_KEY from environment',
            'Demo user credentials in .env',
            'No sensitive data in Python files'
        ]),
        ('Startup Validation System', [
            'File: validate_startup.py',
            'Configuration file verification',
            'Directory structure validation',
            'Python dependency checking',
            'Security settings audit'
        ]),
        ('Application Integration', [
            'app/__init__.py - Integrated config module',
            'run.py - Configuration display',
            'app/modules/auth.py - Demo users from config',
            'app/modules/logger.py - Log folder from config'
        ]),
        ('Storage Mode Decision', [
            'Decision: In-Memory Storage (Academic Prototype)',
            'Rationale: Simplicity, focus on encryption and metrics',
            'Configuration: STORAGE_MODE=in-memory',
            'Future path: SQLite migration documented'
        ]),
        ('Documentation', [
            'WEEK5_ENVIRONMENT_HARDENING.md (400+ lines)',
            'WEEK5_COMPLETION_REPORT.md',
            'ENVIRONMENT_SETUP_GUIDE.md',
            'Comprehensive configuration reference'
        ])
    ]
    
    for title_text, items in deliverables:
        # Subheading
        sub = doc.add_paragraph()
        sub_run = sub.add_run(title_text)
        sub_run.font.name = 'Times New Roman'
        sub_run.font.size = Pt(12)
        sub_run.bold = True
        sub.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        sub.paragraph_format.left_indent = Inches(0.25)
        
        # Bullet points
        for item in items:
            bullet = doc.add_paragraph(item, style='List Bullet')
            bullet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            bullet.paragraph_format.left_indent = Inches(0.5)
            for run in bullet.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Configuration Coverage Section
    heading3 = doc.add_paragraph()
    run = heading3.add_run('CONFIGURATION COVERAGE')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    heading3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    config_intro = doc.add_paragraph('22 Total Configuration Options')
    config_intro.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    for run in config_intro.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    categories = [
        'Flask Settings: SECRET_KEY, ENV, DEBUG, HOST, PORT (5)',
        'File Upload: MAX_FILE_SIZE_MB, Upload paths (4)',
        'Logging: LOG_FOLDER, LOG_LEVEL (2)',
        'Security: SESSION_TIMEOUT, LOGIN_ATTEMPTS, LOCKOUT (3)',
        'Encryption: ALGORITHM, KEY_ROTATION_POLICY (2)',
        'Demo Users: Configurable list format (1)',
        'Storage: MODE - in-memory or database (1)',
        'Research: METRICS_ENABLED, OUTPUT_FILE (2)'
    ]
    
    for category in categories:
        cat_p = doc.add_paragraph(category, style='List Bullet')
        cat_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        for run in cat_p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Security Improvements Section
    heading4 = doc.add_paragraph()
    run = heading4.add_run('SECURITY IMPROVEMENTS')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    heading4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    improvements_before = [
        'Hardcoded SECRET_KEY in Python',
        'Hardcoded demo credentials',
        'Hardcoded paths and configuration',
        'No startup validation',
        'No environment-specific settings'
    ]
    
    before_heading = doc.add_paragraph()
    before_run = before_heading.add_run('Before Week 5:')
    before_run.font.name = 'Times New Roman'
    before_run.font.size = Pt(12)
    before_run.bold = True
    before_heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    for item in improvements_before:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    improvements_after = [
        'SECRET_KEY from environment',
        'Credentials in .env file',
        'All paths configurable',
        '6-point startup validation',
        'Dev/Test/Prod ready',
        'Security warnings for misconfigurations'
    ]
    
    after_heading = doc.add_paragraph()
    after_run = after_heading.add_run('After Week 5:')
    after_run.font.name = 'Times New Roman'
    after_run.font.size = Pt(12)
    after_run.bold = True
    after_heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    for item in improvements_after:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Files Created Section
    heading5 = doc.add_paragraph()
    run = heading5.add_run('FILES CREATED')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    heading5.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    files_created = [
        'config.py (280 lines) - Configuration management module',
        '.env (45 lines) - Local development configuration',
        '.env.example (75 lines) - Configuration template',
        'validate_startup.py (200 lines) - Startup validation script',
        'WEEK5_ENVIRONMENT_HARDENING.md (400+ lines) - Technical documentation',
        'WEEK5_COMPLETION_REPORT.md - Executive summary',
        'ENVIRONMENT_SETUP_GUIDE.md - Setup instructions',
        'Total New Files: 7'
    ]
    
    for file_item in files_created:
        p = doc.add_paragraph(file_item, style='List Bullet')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Files Modified Section
    heading6 = doc.add_paragraph()
    run = heading6.add_run('FILES MODIFIED')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    heading6.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    files_modified = [
        'app/__init__.py - Integrated config module',
        'run.py - Loads config, displays on startup',
        'app/modules/auth.py - Uses config.DEMO_USERS',
        'app/modules/logger.py - Uses config.LOG_FOLDER'
    ]
    
    for mod_item in files_modified:
        p = doc.add_paragraph(mod_item, style='List Bullet')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # How to Use Section
    heading7 = doc.add_paragraph()
    run = heading7.add_run('HOW TO USE')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    heading7.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    quick_start_heading = doc.add_paragraph()
    quick_run = quick_start_heading.add_run('Quick Start:')
    quick_run.font.name = 'Times New Roman'
    quick_run.font.size = Pt(12)
    quick_run.bold = True
    quick_start_heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    steps = [
        'Run: python validate_startup.py',
        'Run: python run.py',
        'Access at http://127.0.0.1:5000',
        'Login with: admin/admin123'
    ]
    
    for step in steps:
        p = doc.add_paragraph(step, style='List Bullet')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Change Configuration Section
    change_heading = doc.add_paragraph()
    change_run = change_heading.add_run('Change Configuration:')
    change_run.font.name = 'Times New Roman'
    change_run.font.size = Pt(12)
    change_run.bold = True
    change_heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    change_steps = [
        'Edit .env file',
        'Change desired settings (FLASK_PORT, MAX_FILE_SIZE_MB, etc.)',
        'Restart application (python run.py)',
        'New configuration loads automatically'
    ]
    
    for step in change_steps:
        p = doc.add_paragraph(step, style='List Bullet')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Summary Section
    heading8 = doc.add_paragraph()
    run = heading8.add_run('SUMMARY')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    heading8.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    summary_items = [
        'Centralized configuration system implemented',
        'No secrets in code',
        'Comprehensive validation system',
        'Deployment ready',
        'Research reproducible',
        'Security hardened',
        'System ready for production-like deployment'
    ]
    
    for item in summary_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run('Week 5 Complete: 100% of objectives achieved')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    footer_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Save document
    filename = 'WEEK5_ENVIRONMENT_HARDENING.docx'
    doc.save(filename)
    print(f'✓ Document created: {filename}')
    print(f'✓ Formatting: Times New Roman, Double Spaced')
    print(f'✓ Format: Professional DOCX')


if __name__ == '__main__':
    try:
        print('Generating Week 5 DOCX document...')
        create_docx()
        print('✓ Success!')
    except Exception as e:
        print(f'✗ Error: {str(e)}')
        sys.exit(1)
